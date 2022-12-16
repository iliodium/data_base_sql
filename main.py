import os
import glob
import uuid
import random
import psycopg2
import numpy as np
import scipy.interpolate
import scipy.io as sio
import matplotlib.cm as cm
import matplotlib.tri as mtri
import matplotlib.pyplot as plt
from matplotlib.colors import Normalize

from docx import Document
from docx.shared import Inches, Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from psycopg2 import Error
from celluloid import Camera
from scipy.fft import fft, rfftfreq
from scipy.signal import argrelextrema, welch


class Artist:
    """Класс отвечает за отрисовку графиков

    Виды графиков:
        -Изополя:
            -Максимальные значения
            -Средние значения
            -Минимальные значения
            -Среднее квадратичное отклонение

        -Огибающие:
            -Максимальные значения
            -Минимальные значения

        -Коэффициенты обеспеченности:
            -Максимальные значения
            -Минимальные значения

        -Спектры

        -Нестационарные сигналы

    """
    @staticmethod
    def disc_isofield(mode, pressure_coefficients, coordinates, alpha, model_name, angle = 0):
        """Отрисовка дискретных изополей"""
        integral_func = []
        mods = {
            'max': np.max(pressure_coefficients, axis=0),
            'mean': np.mean(pressure_coefficients, axis=0),
            'min': np.min(pressure_coefficients, axis=0),
            'std': np.std(pressure_coefficients, axis=0),
        }  # Виды изополей
        pressure_coefficients = mods[mode]
        min_v = np.min(pressure_coefficients)
        max_v = np.max(pressure_coefficients)
        steps = {
            'max': 0.2,
            'mean': 0.2 if alpha == 6 else 0.1,
            'min': 0.2,
            'std': 0.05,
        }  # Шаги для изополей
        #  Шаг для изополей и контурных линий
        levels = np.arange(np.round(min_v, 1), np.round(max_v, 1) + 0.1, steps[mode])
        x, z = np.array(coordinates)
        breadth, depth, height = int(model_name[0]) / 10, int(model_name[1]) / 10, int(model_name[2]) / 10
        count_sensors_on_model = len(pressure_coefficients)
        count_sensors_on_middle = int(model_name[0]) * 5
        count_sensors_on_side = int(model_name[1]) * 5
        count_row = count_sensors_on_model // (2 * (count_sensors_on_middle + count_sensors_on_side))
        pressure_coefficients = np.reshape(pressure_coefficients, (count_row, -1))
        pressure_coefficients = np.split(pressure_coefficients, [count_sensors_on_middle,
                                                                 count_sensors_on_middle + count_sensors_on_side,
                                                                 2 * count_sensors_on_middle + count_sensors_on_side,
                                                                 2 * (
                                                                             count_sensors_on_middle + count_sensors_on_side)
                                                                 ], axis=1)
        x = np.reshape(x, (count_row, -1))
        x = np.split(x, [count_sensors_on_middle,
                         count_sensors_on_middle + count_sensors_on_side,
                         2 * count_sensors_on_middle + count_sensors_on_side,
                         2 * (count_sensors_on_middle + count_sensors_on_side)
                         ], axis=1)

        z = np.reshape(z, (count_row, -1))
        z = np.split(z, [count_sensors_on_middle,
                         count_sensors_on_middle + count_sensors_on_side,
                         2 * count_sensors_on_middle + count_sensors_on_side,
                         2 * (count_sensors_on_middle + count_sensors_on_side)
                         ], axis=1)

        del pressure_coefficients[4]
        del x[4]
        del z[4]

        for i in range(4):
            x[i] = x[i].tolist()
            z[i] = z[i].tolist()

        x1, x2, x3, x4 = x  # x это тензор со всеми координатами граней по ширине,x1...x4 координаты отдельных граней
        z1, z2, z3, z4 = z  # z это тензор со всеми координатами граней по высоте,z1...z4 координаты отдельных граней

        x = [np.array(x1), np.array(x2), np.array(x3), np.array(x4)]  # Входные координаты для изополей
        z = [np.array(z1), np.array(z2), np.array(z3), np.array(z4)]  # Входные координаты для изополей
        ret_int = []  # массив с функциями изополей
        # Расширение матрицы координат по бокам
        for i in range(len(x1)):
            x1[i] = [0] + x1[i] + [breadth]
            x2[i] = [breadth] + x2[i] + [breadth + depth]
            x3[i] = [breadth + depth] + x3[i] + [2 * breadth + depth]
            x4[i] = [2 * breadth + depth] + x4[i] + [2 * (breadth + depth)]

        x1.append(x1[0])
        x2.append(x2[0])
        x3.append(x3[0])
        x4.append(x4[0])

        x1.insert(0, x1[0])
        x2.insert(0, x2[0])
        x3.insert(0, x3[0])
        x4.insert(0, x4[0])

        x_extended = [np.array(x1), np.array(x2), np.array(x3), np.array(x4)]  # Расширенные координаты для изополей

        # Расширение матрицы координат по бокам
        for i in range(len(z1)):
            z1[i] = [z1[i][0]] + z1[i] + [z1[i][0]]
            z2[i] = [z2[i][0]] + z2[i] + [z2[i][0]]
            z3[i] = [z3[i][0]] + z3[i] + [z3[i][0]]
            z4[i] = [z4[i][0]] + z4[i] + [z4[i][0]]

        z1.append([0 for _ in range(len(z1[0]))])
        z2.append([0 for _ in range(len(z2[0]))])
        z3.append([0 for _ in range(len(z3[0]))])
        z4.append([0 for _ in range(len(z4[0]))])

        z1.insert(0, [height for _ in range(len(z1[0]))])
        z2.insert(0, [height for _ in range(len(z2[0]))])
        z3.insert(0, [height for _ in range(len(z3[0]))])
        z4.insert(0, [height for _ in range(len(z4[0]))])

        z_extended = [np.array(z1), np.array(z2), np.array(z3), np.array(z4)]  # Расширенные координаты для изополей
        fig, graph = plt.subplots(1, 4, dpi=200, num=1, clear=True, figsize=(9, 5))
        cmap = cm.get_cmap(name="jet")
        data_colorbar = None
        normalizer = Normalize(min_v, max_v)
        im = cm.ScalarMappable(norm=normalizer, cmap=cmap)
        for i in range(4):

            x_old = x[i].reshape(1, -1)[0]
            # Вычитаем чтобы все координаты по x находились в интервале [0, 1]
            if i == 1:
                x_old -= breadth

            elif i == 2:
                x_old -= (breadth + depth)

            elif i == 3:
                x_old -= (2 * breadth + depth)

            x_old = x_old[:count_sensors_on_middle]
            z_old = z[i].T[0][:count_sensors_on_side]
            data_old = pressure_coefficients[i]

            X, Y = np.meshgrid(x_old, z_old)
            data_colorbar = graph[i].pcolormesh(X, Y, data_old, cmap=cmap)
            graph[i].plot(X, Y, '.k')
            graph[i].set_ylim([0, height])
            if breadth == depth == height:
                graph[i].set_aspect('equal')
            if i in [0, 2]:
                graph[i].set_xlim([0, breadth])
                graph[i].set_xticks(ticks=np.arange(0, breadth + 0.1, 0.1), fontsize=10)
            else:
                graph[i].set_xlim([0, depth])
                graph[i].set_xticks(ticks=np.arange(0, depth + 0.1, 0.1), fontsize=10)

        fig.colorbar(im, ax=graph.ravel().tolist(), location='bottom', cmap=cmap, ticks=levels).ax.tick_params(
            labelsize=7)

        plt.savefig(
            f'{os.getcwd()}\\Отчет {model_name}_{alpha}\\Изополя ветровых нагрузок и воздействий\\Дискретные\\Изополя {model_name}_{alpha}_{angle:02} {mode}',
            bbox_inches='tight')

    @staticmethod
    def interpolator(coords, val):
        return scipy.interpolate.RBFInterpolator(coords, val, kernel='cubic')

    @staticmethod
    def film(coeffs, coordinates, model_name):
        print('START')
        min_v = np.min(coeffs)
        max_v = np.max(coeffs)
        #  Шаг для изополей и контурных линий
        levels = np.arange(np.round(min_v, 1), np.round(max_v, 1) + 0.1, 0.1)
        x, z = np.array(coordinates)

        breadth, depth, height = int(model_name[0]) / 10, int(model_name[1]) / 10, int(model_name[2]) / 10
        count_sensors_on_model = len(coeffs[0])
        count_sensors_on_middle = int(model_name[0]) * 5
        count_sensors_on_side = int(model_name[1]) * 5
        count_row = count_sensors_on_model // (2 * (count_sensors_on_middle + count_sensors_on_side))
        x = np.reshape(x, (count_row, -1))
        x = np.split(x, [count_sensors_on_middle,
                         count_sensors_on_middle + count_sensors_on_side,
                         2 * count_sensors_on_middle + count_sensors_on_side,
                         2 * (count_sensors_on_middle + count_sensors_on_side)
                         ], axis=1)
        z = np.reshape(z, (count_row, -1))
        z = np.split(z, [count_sensors_on_middle,
                         count_sensors_on_middle + count_sensors_on_side,
                         2 * count_sensors_on_middle + count_sensors_on_side,
                         2 * (count_sensors_on_middle + count_sensors_on_side)
                         ], axis=1)

        del x[4]
        del z[4]

        for i in range(4):
            x[i] = x[i].tolist()
            z[i] = z[i].tolist()

        x1, x2, x3, x4 = x  # x это тензор со всеми координатами граней по ширине,x1...x4 координаты отдельных граней
        z1, z2, z3, z4 = z  # z это тензор со всеми координатами граней по высоте,z1...z4 координаты отдельных граней

        x = [np.array(x1), np.array(x2), np.array(x3), np.array(x4)]  # Входные координаты для изополей
        z = [np.array(z1), np.array(z2), np.array(z3), np.array(z4)]  # Входные координаты для изополей
        ret_int = []  # массив с функциями изополей

        # Расширение матрицы координат по бокам
        for i in range(len(x1)):
            x1[i] = [0] + x1[i] + [breadth]
            x2[i] = [breadth] + x2[i] + [breadth + depth]
            x3[i] = [breadth + depth] + x3[i] + [2 * breadth + depth]
            x4[i] = [2 * breadth + depth] + x4[i] + [2 * (breadth + depth)]

        x1.append(x1[0])
        x2.append(x2[0])
        x3.append(x3[0])
        x4.append(x4[0])

        x1.insert(0, x1[0])
        x2.insert(0, x2[0])
        x3.insert(0, x3[0])
        x4.insert(0, x4[0])

        x_extended = [np.array(x1), np.array(x2), np.array(x3), np.array(x4)]  # Расширенные координаты для изополей

        # Расширение матрицы координат по бокам
        for i in range(len(z1)):
            z1[i] = [z1[i][0]] + z1[i] + [z1[i][0]]
            z2[i] = [z2[i][0]] + z2[i] + [z2[i][0]]
            z3[i] = [z3[i][0]] + z3[i] + [z3[i][0]]
            z4[i] = [z4[i][0]] + z4[i] + [z4[i][0]]

        z1.append([0 for _ in range(len(z1[0]))])
        z2.append([0 for _ in range(len(z2[0]))])
        z3.append([0 for _ in range(len(z3[0]))])
        z4.append([0 for _ in range(len(z4[0]))])

        z1.insert(0, [height for _ in range(len(z1[0]))])
        z2.insert(0, [height for _ in range(len(z2[0]))])
        z3.insert(0, [height for _ in range(len(z3[0]))])
        z4.insert(0, [height for _ in range(len(z4[0]))])

        z_extended = [np.array(z1), np.array(z2), np.array(z3), np.array(z4)]  # Расширенные координаты для изополей

        fig, graph = plt.subplots(1, 4, figsize=(16, 9))
        camera = Camera(fig)

        cmap = cm.get_cmap(name="jet")

        x_new = []
        x_old = []
        z_new = []
        z_old = []
        coords = []
        data_new = []
        for i in range(4):
            # x это координаты по ширине
            x_new.append(x_extended[i].reshape(1, -1)[0])
            x_old.append(x[i].reshape(1, -1)[0])
            # Вычитаем чтобы все координаты по x находились в интервале [0, 1]
            if i == 1:
                x_old[i] = x_old[i] - breadth
                x_new[i] = x_new[i] - breadth
            elif i == 2:
                x_old[i] = x_old[i] - (breadth + depth)
                x_new[i] = x_new[i] - (breadth + depth)
            elif i == 3:
                x_old[i] = x_old[i] - (2 * breadth + depth)
                x_new[i] = x_new[i] - (2 * breadth + depth)
            # z это координаты по высоте
            z_new.append(z_extended[i].reshape(1, -1)[0])
            z_old.append(z[i].reshape(1, -1)[0])
            coords.append([[i1, j1] for i1, j1 in zip(x_old[i], z_old[i])])

            graph[i].set_ylim([0, height])
            if breadth == depth == height:
                graph[i].set_aspect('equal')
            if i in [0, 2]:
                graph[i].set_xlim([0, breadth])
                graph[i].set_xticks(ticks=np.arange(0, breadth + 0.1, 0.1))
            else:
                graph[i].set_xlim([0, depth])
                graph[i].set_xticks(ticks=np.arange(0, depth + 0.1, 0.1))

        qqq = 0
        for pressure_coefficients in coeffs[:300]:
            qqq += 1
            print(qqq)

            pressure_coefficients = np.reshape(pressure_coefficients, (count_row, -1))
            pressure_coefficients = np.split(pressure_coefficients,
                                             [count_sensors_on_middle,
                                              count_sensors_on_middle + count_sensors_on_side,
                                              2 * count_sensors_on_middle + count_sensors_on_side,
                                              2 * (
                                                      count_sensors_on_middle + count_sensors_on_side)
                                              ], axis=1)
            del pressure_coefficients[4]

            for i in range(4):
                data_old = pressure_coefficients[i].reshape(1, -1)[0]
                # Интерполятор полученный на основе имеющихся данных
                interpolator = Artist.interpolator(coords[i], data_old)
                # Получаем данные для несуществующих датчиков
                data_new = [float(interpolator([[X, Y]])) for X, Y in zip(x_new[i], z_new[i])]
                triang = mtri.Triangulation(x_new[i], z_new[i])
                refiner = mtri.UniformTriRefiner(triang)
                grid, value = refiner.refine_field(data_new, subdiv=4)
                data_colorbar = graph[i].tricontourf(grid, value, cmap=cmap, levels=levels, extend='both')
                aq = graph[i].tricontour(grid, value, linewidths=1, linestyles='solid', colors='black', levels=levels)
                graph[i].clabel(aq, fontsize=13)
            camera.snap()

        cbar = fig.colorbar(data_colorbar, ax=graph, location='bottom', cmap=cmap, ticks=levels)
        cbar.ax.tick_params(labelsize=5)
        animation = camera.animate(interval=125)
        animation.save('wind.gif')

        plt.clf()
        plt.close()

    @staticmethod
    def isofield(mode, pressure_coefficients, coordinates, alpha, model_name, angle = 0):
        """Отрисовка изополей"""
        integral_func = []
        mods = {
            'max': np.max(pressure_coefficients, axis=0),
            'mean': np.mean(pressure_coefficients, axis=0),
            'min': np.min(pressure_coefficients, axis=0),
            'std': np.std(pressure_coefficients, axis=0),
        }  # Виды изополей
        pressure_coefficients = mods[mode]
        min_v = np.min(pressure_coefficients)
        max_v = np.max(pressure_coefficients)
        steps = {
            'max': 0.2,
            'mean': 0.2 if alpha == 6 else 0.1,
            'min': 0.2,
            'std': 0.05,
        }  # Шаги для изополей
        #  Шаг для изополей и контурных линий
        levels = np.arange(np.round(min_v, 1), np.round(max_v, 1) + 0.1, steps[mode]).round(2)
        x, z = np.array(coordinates)
        # if integral[1] == 1:
        #     # Масштабирование
        #     print(model_name, 1.25)
        #     k = 1.25  # коэффициент масштабирования по высоте
        #     z *= k
        # else:
        #     k = 1
        #     print(model_name, 1)
        breadth, depth, height = int(model_name[0]) / 10, int(model_name[1]) / 10, int(model_name[2]) / 10
        count_sensors_on_model = len(pressure_coefficients)
        count_sensors_on_middle = int(model_name[0]) * 5
        count_sensors_on_side = int(model_name[1]) * 5
        count_row = count_sensors_on_model // (2 * (count_sensors_on_middle + count_sensors_on_side))
        pressure_coefficients = np.reshape(pressure_coefficients, (count_row, -1))
        pressure_coefficients = np.split(pressure_coefficients, [count_sensors_on_middle,
                                                                 count_sensors_on_middle + count_sensors_on_side,
                                                                 2 * count_sensors_on_middle + count_sensors_on_side,
                                                                 2 * (count_sensors_on_middle + count_sensors_on_side)
                                                                 ], axis=1)
        x = np.reshape(x, (count_row, -1))
        x = np.split(x, [count_sensors_on_middle,
                         count_sensors_on_middle + count_sensors_on_side,
                         2 * count_sensors_on_middle + count_sensors_on_side,
                         2 * (count_sensors_on_middle + count_sensors_on_side)
                         ], axis=1)

        z = np.reshape(z, (count_row, -1))
        z = np.split(z, [count_sensors_on_middle,
                         count_sensors_on_middle + count_sensors_on_side,
                         2 * count_sensors_on_middle + count_sensors_on_side,
                         2 * (count_sensors_on_middle + count_sensors_on_side)
                         ], axis=1)

        del pressure_coefficients[4]
        del x[4]
        del z[4]

        for i in range(4):
            x[i] = x[i].tolist()
            z[i] = z[i].tolist()

        x1, x2, x3, x4 = x  # x это тензор со всеми координатами граней по ширине,x1...x4 координаты отдельных граней
        z1, z2, z3, z4 = z  # z это тензор со всеми координатами граней по высоте,z1...z4 координаты отдельных граней

        x = [np.array(x1), np.array(x2), np.array(x3), np.array(x4)]  # Входные координаты для изополей
        z = [np.array(z1), np.array(z2), np.array(z3), np.array(z4)]  # Входные координаты для изополей
        ret_int = []  # массив с функциями изополей
        # Расширение матрицы координат по бокам
        for i in range(len(x1)):
            x1[i] = [0] + x1[i] + [breadth]
            x2[i] = [breadth] + x2[i] + [breadth + depth]
            x3[i] = [breadth + depth] + x3[i] + [2 * breadth + depth]
            x4[i] = [2 * breadth + depth] + x4[i] + [2 * (breadth + depth)]

        x1.append(x1[0])
        x2.append(x2[0])
        x3.append(x3[0])
        x4.append(x4[0])

        x1.insert(0, x1[0])
        x2.insert(0, x2[0])
        x3.insert(0, x3[0])
        x4.insert(0, x4[0])

        x_extended = [np.array(x1), np.array(x2), np.array(x3), np.array(x4)]  # Расширенные координаты для изополей

        # Расширение матрицы координат по бокам
        for i in range(len(z1)):
            z1[i] = [z1[i][0]] + z1[i] + [z1[i][0]]
            z2[i] = [z2[i][0]] + z2[i] + [z2[i][0]]
            z3[i] = [z3[i][0]] + z3[i] + [z3[i][0]]
            z4[i] = [z4[i][0]] + z4[i] + [z4[i][0]]

        z1.append([0 for _ in range(len(z1[0]))])
        z2.append([0 for _ in range(len(z2[0]))])
        z3.append([0 for _ in range(len(z3[0]))])
        z4.append([0 for _ in range(len(z4[0]))])

        z1.insert(0, [height for _ in range(len(z1[0]))])
        z2.insert(0, [height for _ in range(len(z2[0]))])
        z3.insert(0, [height for _ in range(len(z3[0]))])
        z4.insert(0, [height for _ in range(len(z4[0]))])

        z_extended = [np.array(z1), np.array(z2), np.array(z3), np.array(z4)]  # Расширенные координаты для изополей
        fig, graph = plt.subplots(1, 4, dpi=200, num=1, clear=True, figsize=(9, 5))
        cmap = cm.get_cmap(name="jet")
        data_colorbar = None
        # data_old_integer = []  # данные для дискретного интегрирования по осям
        # data_for_3d_model = []  # данные для 3D модели
        for i in range(4):
            # x это координаты по ширине
            x_new = x_extended[i].reshape(1, -1)[0]
            x_old = x[i].reshape(1, -1)[0]
            # Вычитаем чтобы все координаты по x находились в интервале [0, 1]
            if i == 1:
                x_old -= breadth
                x_new -= breadth
            elif i == 2:
                x_old -= (breadth + depth)
                x_new -= (breadth + depth)
            elif i == 3:
                x_old -= (2 * breadth + depth)
                x_new -= (2 * breadth + depth)
            # z это координаты по высоте
            z_new = z_extended[i].reshape(1, -1)[0]
            z_old = z[i].reshape(1, -1)[0]

            data_old = pressure_coefficients[i].reshape(1, -1)[0]
            # data_old_integer.append(data_old)
            coords = [[i1, j1] for i1, j1 in zip(x_old, z_old)]  # Старые координаты
            # Интерполятор полученный на основе имеющихся данных
            interpolator = Artist.interpolator(coords, data_old)
            integral_func.append(interpolator)
            # Получаем данные для несуществующих датчиков
            data_new = [float(interpolator([[X, Y]])) for X, Y in zip(x_new, z_new)]

            triang = mtri.Triangulation(x_new, z_new)
            refiner = mtri.UniformTriRefiner(triang)
            grid, value = refiner.refine_field(data_new, subdiv=4)
            # data_for_3d_model.append((grid, value))
            data_colorbar = graph[i].tricontourf(grid, value, cmap=cmap, levels=levels, extend='both')
            aq = graph[i].tricontour(grid, value, linewidths=1, linestyles='solid', colors='black', levels=levels)
            X, Y = np.meshgrid(x_old, z_old)
            graph[i].plot(X, Y, '.k', **dict(markersize=3.7))
            graph[i].clabel(aq, fontsize=10)

            graph[i].set_ylim([0, height])
            if breadth == depth == height:
                graph[i].set_aspect('equal')
            if i in [0, 2]:
                graph[i].set_xlim([0, breadth])
                graph[i].set_xticks(ticks=np.arange(0, breadth + 0.1, 0.1), fontsize=10)
            else:
                graph[i].set_xlim([0, depth])
                graph[i].set_xticks(ticks=np.arange(0, depth + 0.1, 0.1), fontsize=10)
            ret_int.append(interpolator)
            # graph[i].axis('off')
        fig.colorbar(data_colorbar, ax=graph, location='bottom', cmap=cmap, ticks=levels).ax.tick_params(labelsize=7)

        # plt.savefig(f'{folder_out}\\Изополя {model_name}_{alpha} {mode}')
        plt.savefig(
            f'{os.getcwd()}\\Отчет {model_name}_{alpha}\\Изополя ветровых нагрузок и воздействий\\Непрерывные\\Изополя {model_name}_{alpha}_{angle:02} {mode}',
            bbox_inches='tight')
        # plt.show()

        # isofield_model().show(data_for_3d_model, levels)
        # return ret_int
        # # Численное интегрирование
        #
        # if integral[0] == -1:
        #     return None
        # elif integral[0] == 0:
        #     count_zone = count_row
        # else:
        #     count_zone = integral[0]
        # ran = random.uniform
        # face1, face2, face3, face4 = [], [], [], []
        # model = face1, face2, face3, face4
        # count_point = integral[1]
        # step_floor = height / count_row
        # for i in range(4):
        #     top = step_floor
        #     down = 0
        #     for _ in range(count_zone):
        #         floor = []
        #         if i in [0, 2]:
        #             for _ in range(count_point):
        #                 floor.append(integral_func[i]([[ran(0, breadth), ran(down, top)]]))
        #         else:
        #             for _ in range(count_point):
        #                 floor.append(integral_func[i]([[ran(0, depth), ran(down, top)]]))
        #         down = top
        #         top += step_floor
        #         # Обычный метод Монте-Синякина
        #         if i in [0, 2]:
        #             model[i].append(sum(floor) * breadth / count_point)
        #         else:
        #             model[i].append(sum(floor) * depth / count_point)
        # print(face1)
        # print(face2)
        # print(face3)
        # print(face4)

    @staticmethod
    def func(x):
        return 15 * x ** 3 + 21 * x ** 2 + 41 * x + 3 * np.sin(x) * np.cos(x)

    @staticmethod
    def check():
        arr = []
        a = 0
        b = 5
        n = 2000000
        for i in range(n):
            arr.append(Artist.func(random.uniform(a, b)))

        print(sum(arr) * (b - a) / n)

    @staticmethod
    def signal(pressure_coefficients, pressure_coefficients1, alpha, model_name, angle):
        time = [i / 1000 for i in range(5001)]
        plt.plot(time, [i[0] for i in pressure_coefficients[:5001]], label='113')
        plt.plot(time, [i[0] for i in pressure_coefficients1[:5001]], label='115')
        plt.legend()
        plt.show()

    @staticmethod
    def signal1(pressure_coefficients, pressure_coefficients_2):
        time = [i / 1000 for i in range(len(pressure_coefficients))]
        plt.plot(time, [i[206] for i in pressure_coefficients], label='old')
        plt.plot(time, pressure_coefficients_2, label='new')
        plt.legend()
        plt.show()

    @staticmethod
    def spectrum(pressure_coefficients, pressure_coefficients1, alpha, model_name, angle):
        N = len(pressure_coefficients)
        yf = (1 / N) * (np.abs(fft([i[0] for i in pressure_coefficients])))[1:N // 2]
        yf1 = (1 / N) * (np.abs(fft([i[0] for i in pressure_coefficients1])))[1:N // 2]
        FD = 1000
        xf = rfftfreq(N, 1 / FD)[1:N // 2]
        xf = xf[1:]
        yf = yf[1:]
        yf1 = yf1[1:]
        plt.plot(xf[:200], yf[:200], antialiased=True, label='113')
        plt.plot(xf[:200], yf1[:200], antialiased=True, label='115')
        plt.legend()
        plt.show()


class Controller:
    """
    Через класс происходит управление базой данных:
        -создание таблиц
        -заполнение таблиц
        -генерация несуществующих вариантов
    """
    __connection = None
    __path_database = None
    __path_database = 'D:\Projects\mat_to_csv\mat files'
    __extrapolatedAnglesInfoList = {}  # ключ вида T<model_name>_<alpha>_<angle>, значения (коэффициенты, x, z)

    def __init__(self):
        self.cursor = None

    def connect(self, database = '', password = '', user = "postgres", host = "127.0.0.1", port = "5432"):
        """Метод подключается от PostgreSQL"""
        try:
            self.__connection = psycopg2.connect(user=user,
                                                 password=password,
                                                 host=host,
                                                 port=port,
                                                 database=database)
            self.cursor = self.__connection.cursor()
            print("Соединение с PostgreSQL открыто")
        except psycopg2.OperationalError:
            assert False, ("Проверьте правильность веденных данных")
        except Exception as error:
            assert False, (f"Ошибка - {error}")

    @staticmethod
    def read_mat(path):
        """Считывает и обрабатывает данные из .mat файла"""
        alpha = path[-11:-10]
        mat_file = sio.loadmat(path)
        breadth = float(mat_file['Building_breadth'][0][0])
        depth = float(mat_file['Building_depth'][0][0])
        height = float(mat_file['Building_height'][0][0])
        model_name = int(f'{int(breadth * 10)}{int(depth * 10)}{int(height * 10)}')
        frequency = int(mat_file['Sample_frequency'][0][0])
        period = float(mat_file['Sample_period'][0][0])
        speed = float(mat_file['Uh_AverageWindSpeed'][0])
        x = [float('%.5f' % i) for i in mat_file["Location_of_measured_points"][0]]
        z = [float('%.5f' % i) for i in mat_file["Location_of_measured_points"][1]]
        sensor_number = [int(i) for i in mat_file["Location_of_measured_points"][2]]
        face_number = [int(i) for i in mat_file["Location_of_measured_points"][3]]
        angle = int(mat_file['Wind_direction_angle'][0][0])
        pressure_coefficients = mat_file["Wind_pressure_coefficients"] * 1000
        pressure_coefficients = pressure_coefficients.round(0)
        pressure_coefficients = pressure_coefficients.astype('int32')
        pressure_coefficients = pressure_coefficients.tolist()
        return {'alpha': alpha,
                'model_name': model_name,
                'breadth': breadth,
                'depth': depth,
                'height': height,
                'frequency': frequency,
                'period': period,
                'speed': speed,
                'x': x,
                'z': z,
                'sensor_number': sensor_number,
                'face_number': face_number,
                'angle': angle,
                'pressure_coefficients': pressure_coefficients,
                }

    @staticmethod
    def converter_coordinates(x_old, depth, breadth, face_number, count_sensors):
        """Возвращает из (x_old) -> (x,y)"""
        x = []
        y = []
        for i in range(count_sensors):
            if face_number[i] == 1:
                x.append(float('%.5f' % (-depth / 2)))
                y.append(float('%.5f' % (breadth / 2 - x_old[i])))
            elif face_number[i] == 2:
                x.append(float('%.5f' % (- depth / 2 + x_old[i] - breadth)))
                y.append(float('%.5f' % (-breadth / 2)))
            elif face_number[i] == 3:
                x.append(float('%.5f' % (depth / 2)))
                y.append(float('%.5f' % (-3 * breadth / 2 + x_old[i] - depth)))
            else:
                x.append(float('%.5f' % (3 * depth / 2 - x_old[i] + 2 * breadth)))
                y.append(float('%.5f' % (breadth / 2)))

        return x, y

    def fill_models_alpha(self, parameters):
        """Добавляет запись в таблицу models_alpha_<alpha>"""
        alpha = parameters['alpha']
        model_name = parameters['model_name']
        angle = parameters['angle']
        pressure_coefficients = parameters['pressure_coefficients']
        flag = 0
        try:
            name = f"T{model_name}_{alpha}_{angle:03d}_1.mat"
            if alpha == '6':
                self.cursor.execute("""
                           select model_id 
                           from experiments_alpha_6
                           where model_name = (%s)
                       """, (model_name,))
            elif alpha == '4':
                self.cursor.execute("""
                           select model_id 
                           from experiments_alpha_4
                           where model_name = (%s)
                       """, (model_name,))

            model_id = self.cursor.fetchall()[0][0]
            if alpha == '6':
                self.cursor.execute("""
                                       select model_id 
                                       from models_alpha_6
                                       where model_id = (%s) and angle = (%s)
                                        """, (model_id, angle))
            elif alpha == '4':
                self.cursor.execute("""
                           select model_id 
                           from models_alpha_4
                           where model_id = (%s) and angle = (%s)
                            """, (model_id, angle))
            check_exists = self.cursor.fetchall()
            if check_exists:
                print(f'{name} была ранее добавлена в models_alpha_{alpha}')
                flag = 1

            if flag == 0:
                try:
                    if alpha == '6':
                        self.cursor.execute("""
                                   insert into models_alpha_6
                                   values
                                   ((%s), (%s), (%s))
                               """, (model_id, angle, pressure_coefficients))
                    elif alpha == '4':
                        self.cursor.execute("""
                                   insert into models_alpha_4
                                   values
                                   ((%s), (%s), (%s))
                               """, (model_id, angle, pressure_coefficients))
                    self.__connection.commit()
                    print(f'{name} добавлена в models_alpha_{alpha}')
                except (Exception, Error) as error:
                    print(f"Ошибка при работе с PostgreSQL, файл {model_name}", error)
        except (Exception, Error) as error:
            print(f"Ошибка при работе с PostgreSQL, файл {model_name}", error)

    def fill_experiments_alpha(self, parameters):
        """Добавляет запись в таблицу experiments_alpha_<alpha>"""
        alpha = parameters['alpha']
        angle = parameters['angle']
        model_name = parameters['model_name']
        breadth = parameters['breadth']
        depth = parameters['depth']
        height = parameters['height']
        frequency = parameters['frequency']
        period = parameters['period']
        speed = parameters['speed']
        x = parameters['x']
        z = parameters['z']
        sensor_number = parameters['sensor_number']
        face_number = parameters['face_number']
        flag = 0  # флаг для проверки наличия записи в бд
        name = f"T{model_name}_{alpha}_{angle:03d}_1.mat"
        try:
            if alpha == '6':
                self.cursor.execute("""
                           select model_name 
                           from experiments_alpha_6
                       """)
            elif alpha == '4':
                self.cursor.execute("""
                           select model_name 
                           from experiments_alpha_4
                       """)
            table = self.cursor.fetchall()
            for row in table:
                if row[0] == model_name:
                    print(f'{name} была ранее добавлена в experiments_alpha_{alpha}')
                    flag = 1
                    break
            if flag == 0:
                try:
                    if alpha == '6':
                        self.cursor.execute("""
                                   insert into experiments_alpha_6 (model_name,
                                                                    breadth,
                                                                    depth,
                                                                    height,
                                                                    sample_frequency,
                                                                    sample_period,
                                                                    uh_AverageWindSpeed,
                                                                    x_coordinates,
                                                                    z_coordinates,
                                                                    sensor_number,
                                                                    face_number)
                                   values((%s),(%s),(%s),(%s),(%s),(%s),(%s),(%s),(%s),(%s),(%s))
                               """, (
                            model_name, breadth, depth, height, frequency, period, speed, x, z, sensor_number,
                            face_number))
                    elif alpha == '4':
                        self.cursor.execute(f"""
                                   insert into experiments_alpha_4 (model_name,
                                                                    breadth,
                                                                    depth,
                                                                    height,
                                                                    sample_frequency,
                                                                    sample_period,
                                                                    uh_AverageWindSpeed,
                                                                    x_coordinates,
                                                                    z_coordinates,
                                                                    sensor_number,
                                                                    face_number)
                                   values((%s),(%s),(%s),(%s),(%s),(%s),(%s),(%s),(%s),(%s),(%s))
                               """, (
                            model_name, breadth, depth, height, frequency, period, speed, x, z, sensor_number,
                            face_number))
                    self.__connection.commit()
                    print(f'{name} добавлена в experiments_alpha_{alpha}')
                except (Exception, Error) as error:
                    print(f"Ошибка при работе с PostgreSQL, файл {model_name}", error)
        except (Exception, Error) as error:
            print(f"Ошибка при работе с PostgreSQL, файл {model_name}", error)

    def fill_db(self, path = None):
        """
        Метод заполняет таблицу experiments_alpha_<alpha> и models_alpha_<alpha> данными из .mat файла.

        Если путь не передан, то в таблицы добавляются все данные из директории.
        """

        if path:
            parameters = self.read_mat(path)
            self.fill_experiments_alpha(parameters)
            self.fill_models_alpha(parameters)

        else:
            self.check_path()
            for alpha in os.listdir(self.__path_database):
                for files in os.listdir(f'{self.__path_database}\\{alpha}'):
                    for file in os.listdir(f"{self.__path_database}\\{alpha}\\{files}"):
                        self.fill_db(f"{self.__path_database}\\{alpha}\\{files}\\{file}")

    def get_paths(self):
        """Возвращает список путей всех файлов"""
        self.paths = []
        self.check_path()
        for alpha in os.listdir(self.__path_database):
            for files in os.listdir(f'{self.__path_database}\\{alpha}'):
                for file in os.listdir(f"{self.__path_database}\\{alpha}\\{files}"):
                    self.paths.append(f"{self.__path_database}\\{alpha}\\{files}\\{file}")

        return self.paths

    def check_path(self):
        """Метод проверяет наличие заданного пути"""
        if self.__path_database is None:
            self.__path_database = input('Введите полный путь до данных:\n')

    def disconnect(self):
        """Метод отключается от PostgreSQL"""
        if self.__connection:
            self.cursor.close()
            self.__connection.close()
            print("Соединение с PostgreSQL закрыто")

    def create_tables(self):
        """Метод создаёт таблицы"""
        # создается таблица с экспериментами
        try:
            self.cursor.execute('''
                create table if not exists experiments_alpha_4
                (
                    model_id serial primary key,
                    model_name smallint not null,
                    breadth real not null,
                    depth real not null,
                    height real not null,
                    sample_frequency smallint not null,
                    sample_period real not null,
                    uh_AverageWindSpeed real not null,
                    x_coordinates real[] not null,
                    z_coordinates real[] not null,
                    sensor_number smallint[] not null,
                    face_number smallint[] not null
                )''')
            self.__connection.commit()
            print("Создана таблица с экспериментами с параметром 4")
        except Exception as error:
            print(f"Ошибка -> {error}")
            self.__connection.commit()

        try:
            self.cursor.execute('''
                create table if not exists experiments_alpha_6
                (
                    model_id serial primary key,
                    model_name smallint not null,
                    breadth real not null,
                    depth real not null,
                    height real not null,
                    sample_frequency smallint not null,
                    sample_period real not null,
                    uh_AverageWindSpeed real not null,
                    x_coordinates real[] not null,
                    z_coordinates real[] not null,
                    sensor_number smallint[] not null,
                    face_number smallint[] not null
                )''')
            self.__connection.commit()
            print("Создана таблица с экспериментами с параметром 6")
        except Exception as error:
            print(f"Ошибка -> {error}")
            self.__connection.commit()

        # создается таблица с моделями
        try:
            self.cursor.execute(f"""
                create table if not exists models_alpha_4
                (
                    model_id serial not null,
                    angle smallint not null,
                    pressure_coefficients smallint[][] not null,
                    
                    constraint FK_{str(uuid.uuid4()).replace('-', '')} foreign key (model_id) 
                    references experiments_alpha_4(model_id)
                )""")
            self.__connection.commit()
            print(f"Создана таблица с моделями с параметром 4")
        except Exception as error:
            print(f"Ошибка -> {error}")
            self.__connection.commit()

        try:
            self.cursor.execute(f"""
                create table if not exists models_alpha_6
                (
                    model_id serial not null,
                    angle smallint not null,
                    pressure_coefficients smallint[][] not null,

                    constraint FK_{str(uuid.uuid4()).replace('-', '')} foreign key (model_id) 
                    references experiments_alpha_6(model_id)
                )""")
            self.__connection.commit()
            print(f"Создана таблица с моделями с параметром 6")

        except Exception as error:
            print(f"Ошибка -> {error}")
            self.__connection.commit()

    def generate_not_exists_case(self, alpha, model_name, angle):
        """Определяет параметры для генерация несуществующего варианта"""
        print(f'Генерация модели {model_name} с параметром {alpha} и углом {angle}')
        if model_name[0] == model_name[1]:  # в основании квадрат
            if any([45 < angle < 90, 135 < angle < 180, 225 < angle < 270, 315 < angle < 360]):
                self.__extrapolatedAnglesInfoList[f'T{model_name}_{alpha}_{angle:03d}'] = self.generator('rev',
                                                                                                         alpha,
                                                                                                         model_name,
                                                                                                         angle)
            else:
                self.__extrapolatedAnglesInfoList[f'T{model_name}_{alpha}_{angle:03d}'] = self.generator('for',
                                                                                                         alpha,
                                                                                                         model_name,
                                                                                                         angle)
        else:  # в основании прямоугольник
            if any([90 < angle < 180, 270 < angle < 360]):
                self.__extrapolatedAnglesInfoList[f'T{model_name}_{alpha}_{angle:03d}'] = self.generator('rev',
                                                                                                         alpha,
                                                                                                         model_name,
                                                                                                         angle)
            else:
                self.__extrapolatedAnglesInfoList[f'T{model_name}_{alpha}_{angle:03d}'] = self.generator('for',
                                                                                                         alpha,
                                                                                                         model_name,
                                                                                                         angle)

    def generator(self, type_gen, alpha, model_name, angle):
        """Определяет как менять расстановку датчиков и вызывает расстановщика"""
        angle_parent = {
            'rev': 90 * (angle // 90 + 1) - angle,
            'for': angle % 90
        }

        # определение параметров для перестановки данных
        type_view = {
            True: {
                'rev': {
                    45 < angle < 90: (1, 0, 3, 2),
                    135 < angle < 180: (2, 1, 0, 3),
                    225 < angle < 270: (3, 2, 1, 0),
                    315 < angle < 360: (0, 3, 1, 2)
                },
                'for': {
                    0 <= angle <= 45: (0, 1, 2, 3),
                    90 <= angle <= 135: (3, 0, 1, 2),
                    180 <= angle <= 225: (2, 3, 0, 1),
                    270 <= angle <= 315: (1, 2, 3, 0)
                }},
            False: {
                'rev': {
                    90 < angle < 180: (2, 1, 0, 3),
                    270 < angle < 360: (0, 3, 2, 1),
                },
                'for': {
                    0 <= angle <= 90: (0, 1, 2, 3),
                    180 <= angle <= 270: (2, 3, 0, 1)
                }}
        }
        # model_name[0] == model_name[1] проверка на квадрат в основании
        view_data = type_view[model_name[0] == model_name[1]][type_gen][True]
        return self.changer_sequence_sensors(alpha, model_name, angle_parent[type_gen], view_data, type_gen)

    def changer_sequence_sensors(self, alpha, model_name, angle, view_data = (0, 1, 2, 3), type_gen = 'for'):
        """Меняет порядок следования датчиков и координат тем самым генерируя не существующие углы"""
        coeffs = self.get_pressure_coefficients(alpha, model_name, angle)
        x, z = self.get_coordinates(alpha, model_name)

        f1, f2, f3, f4 = view_data
        count_sensors_on_middle = int(model_name[0]) * 5
        count_sensors_on_side = int(model_name[1]) * 5
        count_sensors_on_model = len(coeffs[0])
        count_row = count_sensors_on_model // (2 * (count_sensors_on_middle + count_sensors_on_side))

        if type_gen == 'for':
            for i in range(len(coeffs)):
                arr = np.array(coeffs[i]).reshape((count_row, -1))
                arr = np.split(arr, [count_sensors_on_middle,
                                     count_sensors_on_middle + count_sensors_on_side,
                                     2 * count_sensors_on_middle + count_sensors_on_side,
                                     2 * (count_sensors_on_middle + count_sensors_on_side)
                                     ], axis=1)

                coeffs[i] = np.concatenate((arr[f1],
                                            arr[f2],
                                            arr[f3],
                                            arr[f4]), axis=1).reshape(count_sensors_on_model)
            arr = np.array(x).reshape((count_row, -1))
            arr = np.split(arr, [count_sensors_on_middle,
                                 count_sensors_on_middle + count_sensors_on_side,
                                 2 * count_sensors_on_middle + count_sensors_on_side,
                                 2 * (count_sensors_on_middle + count_sensors_on_side)
                                 ], axis=1)

            x = np.concatenate((arr[f1],
                                arr[f2],
                                arr[f3],
                                arr[f4]), axis=1).reshape(count_sensors_on_model)

        else:
            for i in range(len(coeffs)):
                arr = np.array(coeffs[i]).reshape((count_row, -1))
                arr = np.split(arr, [count_sensors_on_middle,
                                     count_sensors_on_middle + count_sensors_on_side,
                                     2 * count_sensors_on_middle + count_sensors_on_side,
                                     2 * (count_sensors_on_middle + count_sensors_on_side)
                                     ], axis=1)

                coeffs[i] = np.concatenate((np.flip(arr[f1], axis=1),
                                            np.flip(arr[f2], axis=1),
                                            np.flip(arr[f3], axis=1),
                                            np.flip(arr[f4], axis=1)), axis=1).reshape(count_sensors_on_model)
            arr = np.array(x).reshape((count_row, -1))
            arr = np.split(arr, [count_sensors_on_middle,
                                 count_sensors_on_middle + count_sensors_on_side,
                                 2 * count_sensors_on_middle + count_sensors_on_side,
                                 2 * (count_sensors_on_middle + count_sensors_on_side)
                                 ], axis=1)

            x = np.concatenate((np.flip(arr[f1], axis=1),
                                np.flip(arr[f2], axis=1),
                                np.flip(arr[f3], axis=1),
                                np.flip(arr[f4], axis=1)), axis=1).reshape(count_sensors_on_model)
        return coeffs, x, z

    def get_coordinates(self, alpha, model_name):
        """Возвращает координаты датчиков"""
        if alpha == '4' or alpha == 4:
            self.cursor.execute("""
                        select x_coordinates, z_coordinates
                        from experiments_alpha_4
                        where model_name = (%s)
                    """, (model_name,))

        elif alpha == '6' or alpha == 6:
            self.cursor.execute("""
                        select x_coordinates, z_coordinates
                        from experiments_alpha_6
                        where model_name = (%s)
                    """, (model_name,))
        self.__connection.commit()
        x, z = self.cursor.fetchall()[0]
        return x, z

    def get_pressure_coefficients(self, alpha, model_name, angle):
        """Возвращает коэффициенты давления"""
        if alpha == '4' or alpha == 4:
            self.cursor.execute("""
                        select pressure_coefficients
                        from models_alpha_4
                        where model_id = (
                        select model_id
                        from experiments_alpha_4
                        where model_name = (%s)
                        ) and angle = (%s)
                    """, (model_name, angle))

        elif alpha == '6' or alpha == 6:
            self.cursor.execute("""
                        select pressure_coefficients
                        from models_alpha_6
                        where model_id = (
                        select model_id
                        from experiments_alpha_6
                        where model_name = (%s)
                        ) and angle = (%s)
                    """, (model_name, angle))
        self.__connection.commit()
        pressure_coefficients = self.cursor.fetchall()
        if not pressure_coefficients:
            if model_name[0] == model_name[1] and angle <= 45 or model_name[0] != model_name[1] and angle <= 90:
                print(f"T{model_name}_{alpha}_{angle:03d} не существует и не может быть сгенерирован")
                return None
            else:
                self.generate_not_exists_case(alpha, model_name, angle)
                return self.__extrapolatedAnglesInfoList[f'T{model_name}_{alpha}_{angle:03d}'][0]
        return pressure_coefficients[0][0]

    def graphs(self, mode, alpha, model_name, angle):
        """Метод отвечает за вызов графика из класса Artist"""
        angle = int(angle) % 360
        if angle % 5 != 0:
            print('Углы должны быть кратны 5')
            return None
        pressure_coefficients = np.array(self.get_pressure_coefficients(alpha, model_name, angle)) / 1000
        try:
            coordinates = self.__extrapolatedAnglesInfoList[f'T{model_name}_{alpha}_{angle:03d}'][1:]
        except:
            coordinates = self.get_coordinates(alpha, model_name)

        return Artist.isofield(mode, pressure_coefficients, coordinates, alpha, model_name, angle)

    def generate_signal(self, alpha, model_name, angle, x_gen, z_gen):
        new_signal = []
        pressure_coefficients = np.array(self.get_pressure_coefficients(alpha, model_name, angle)) / 1000
        x, z = self.get_coordinates(alpha, model_name)
        breadth, depth, height = int(model_name[0]) / 10, int(model_name[1]) / 10, int(model_name[2]) / 10
        count_sensors_on_model = len(pressure_coefficients[0])
        count_sensors_on_middle = int(model_name[0]) * 5
        count_sensors_on_side = int(model_name[1]) * 5
        count_row = count_sensors_on_model // (2 * (count_sensors_on_middle + count_sensors_on_side))
        face = -1
        if 0 <= x_gen <= breadth:
            face = 0
        elif breadth < x_gen <= (breadth + depth):
            face = 1
            x_gen -= breadth

        elif (breadth + depth) < x_gen <= (2 * breadth + depth):
            face = 2
            x_gen -= (breadth + depth)

        else:
            face = 3
            x_gen -= (2 * breadth + depth)

        x = np.reshape(x, (count_row, -1))
        x = np.split(x, [count_sensors_on_middle,
                         count_sensors_on_middle + count_sensors_on_side,
                         2 * count_sensors_on_middle + count_sensors_on_side,
                         2 * (count_sensors_on_middle + count_sensors_on_side)
                         ], axis=1)

        z = np.reshape(z, (count_row, -1))
        z = np.split(z, [count_sensors_on_middle,
                         count_sensors_on_middle + count_sensors_on_side,
                         2 * count_sensors_on_middle + count_sensors_on_side,
                         2 * (count_sensors_on_middle + count_sensors_on_side)
                         ], axis=1)

        del x[4]
        del z[4]

        for i in range(4):
            x[i] = x[i].tolist()
            z[i] = z[i].tolist()

        # x1...x4 координаты отдельных граней
        x1, x2, x3, x4 = x  # x это тензор со всеми координатами граней по ширине
        # z1...z4 координаты отдельных граней
        z1, z2, z3, z4 = z  # z это тензор со всеми координатами граней по высоте

        x = [np.array(x1), np.array(x2), np.array(x3), np.array(x4)]  # Входные координаты для изополей
        z = [np.array(z1), np.array(z2), np.array(z3), np.array(z4)]  # Входные координаты для изополей

        # x это координаты по ширине
        x_old = x[face].reshape(1, -1)[0]
        # Вычитаем чтобы все координаты по x находились в интервале [0, 1]
        if face == 1:
            x_old -= breadth
        elif face == 2:
            x_old -= (breadth + depth)
        elif face == 3:
            x_old -= (2 * breadth + depth)
        # z это координаты по высоте
        z_old = z[face].reshape(1, -1)[0]
        coords = [[i1, j1] for i1, j1 in zip(x_old, z_old)]  # Старые координаты

        for ind in range(len(pressure_coefficients)):
            coefficient = pressure_coefficients[ind]

            coefficient = np.reshape(coefficient, (count_row, -1))
            coefficient = np.split(coefficient, [count_sensors_on_middle,
                                                 count_sensors_on_middle + count_sensors_on_side,
                                                 2 * count_sensors_on_middle + count_sensors_on_side,
                                                 2 * (
                                                         count_sensors_on_middle + count_sensors_on_side)
                                                 ], axis=1)
            del coefficient[4]
            data_old = coefficient[face].reshape(1, -1)[0]
            # Интерполятор полученный на основе имеющихся данных
            interpolator = Artist.interpolator(coords, data_old)
            new_signal.append(float(interpolator([[x_gen, z_gen]])))

        return new_signal

    def graphic_mean(self, alpha, model_name, angle):
        pr_norm = np.array(self.get_pressure_coefficients(alpha, model_name, angle)) / 1000
        count_sensors_on_model = len(pr_norm[0])
        count_sensors_on_middle = int(model_name[0]) * 5
        count_sensors_on_side = int(model_name[1]) * 5
        count_row = count_sensors_on_model // (2 * (count_sensors_on_middle + count_sensors_on_side))

        sum_int_x = []
        sum_int_y = []

        for coeff in pr_norm:
            coeff = np.reshape(coeff, (count_row, -1))
            coeff = np.split(coeff, [count_sensors_on_middle,
                                     count_sensors_on_middle + count_sensors_on_side,
                                     2 * count_sensors_on_middle + count_sensors_on_side,
                                     2 * (count_sensors_on_middle + count_sensors_on_side)
                                     ], axis=1)
            del coeff[4]
            faces_x = []
            faces_y = []
            for face in range(len(coeff)):
                if face in [0, 2]:
                    faces_x.append(np.sum(coeff[face]) / (count_sensors_on_model / 4))
                else:
                    faces_y.append(np.sum(coeff[face]) / (count_sensors_on_model / 4))

            sum_int_x.append((faces_x[0] - faces_x[1]))
            sum_int_y.append((faces_y[0] - faces_y[1]))
        print(f"""
                среднее по Cy
                {sum(sum_int_y) / 32768}
                среднее по Cx
                {sum(sum_int_x) / 32768}
                """)
        print(f"""
                стандартное откл по Cx
                {np.std(sum_int_x)}
                стандартное откл по Cy
                {np.std(sum_int_y)}
                """)
        fig, graph = plt.subplots(1, 2, figsize=(16, 9))
        graph[0].plot(list(range(1, 32769)), sum_int_x, label='Cx')
        graph[1].plot(list(range(1, 32769)), sum_int_y, label='Cy')
        for i in range(2):
            graph[i].legend()
            graph[i].grid()
        plt.savefig(f'{folder_out}\\график средних {model_name}_{alpha}')
        plt.clf()
        plt.close()

    def spectr_mean(self, alpha, model_name, angle = 0, border = 30000):
        pr_norm = np.array(self.get_pressure_coefficients(alpha, model_name, angle)) / 1000
        count_sensors_on_model = len(pr_norm[0])
        count_sensors_on_middle = int(model_name[0]) * 5
        count_sensors_on_side = int(model_name[1]) * 5
        count_row = count_sensors_on_model // (2 * (count_sensors_on_middle + count_sensors_on_side))

        sum_int_x = []
        sum_int_y = []

        for coeff in pr_norm:
            coeff = np.reshape(coeff, (count_row, -1))
            coeff = np.split(coeff, [count_sensors_on_middle,
                                     count_sensors_on_middle + count_sensors_on_side,
                                     2 * count_sensors_on_middle + count_sensors_on_side,
                                     2 * (count_sensors_on_middle + count_sensors_on_side)
                                     ], axis=1)
            del coeff[4]
            faces_x = []
            faces_y = []
            for face in range(len(coeff)):
                if face in [0, 2]:
                    faces_x.append(np.sum(coeff[face]) / (count_sensors_on_model / 4))
                else:
                    faces_y.append(np.sum(coeff[face]) / (count_sensors_on_model / 4))

            sum_int_x.append((faces_x[0] - faces_x[1]))
            sum_int_y.append((faces_y[0] - faces_y[1]))
        N = len(sum_int_x)

        yfCx = (1 / N) * (np.abs(fft(sum_int_x)))[1:N // 2]
        yfCy = (1 / N) * (np.abs(fft(sum_int_y)))[1:N // 2]

        FD = 1000
        xf = rfftfreq(N, 1 / FD)[1:N // 2]

        fig, graph = plt.subplots(1, 2, figsize=(16, 9))
        if border == -1:
            border = 100000

        graph[0].plot(xf[:border], yfCx[:border], antialiased=True, label='Cx')
        graph[1].plot(xf[:border], yfCy[:border], antialiased=True, label='Cy')
        for i in range(2):
            graph[i].legend()
            graph[i].grid()

        plt.savefig(f'{folder_out}\\спектр средних {model_name}_{alpha}')
        plt.clf()
        plt.close()

    def generate_model_pic(self, alpha, model_name):
        breadth, depth, height = int(model_name[0]) / 10, int(model_name[1]) / 10, int(model_name[2]) / 10
        size_x = 2 * (breadth + depth)
        x, z = self.get_coordinates(alpha, model_name)
        pr_coeff = np.array(self.get_pressure_coefficients(alpha, model_name, 0)) / 1000
        count_sensors = len(pr_coeff[0])
        fig, ax = plt.subplots(figsize=(6.49, 5), dpi=200, num=1, clear=True)
        ax.set_title('Позиция датчиков', fontweight='semibold', fontsize=8)
        ax.set_xlabel('Горизонталь /м', fontweight='semibold', fontsize=8)
        ax.set_ylabel('Вертикаль /м', fontweight='semibold', fontsize=8)
        ax.set_ylim(0, height)
        ax.set_xlim(0, size_x)
        xtick_s = 0.05
        ytick_s = 0.02 if height in [0.1, 0.2] else 0.05
        xticks = np.arange(0, size_x + xtick_s, xtick_s)
        yticks = np.arange(0, height + ytick_s, ytick_s)
        xlabels = ['0'] + [str(i)[:4].rstrip('0') for i in xticks[1:]]
        ylabels = ['0'] + [str(i)[:4].rstrip('0') for i in yticks[1:]]
        ax.set_xticks(xticks, labels=xlabels)
        ax.set_yticks(yticks, labels=ylabels)
        xticks_minor = np.arange(0, size_x, 0.02)
        ax.set_xticks(xticks_minor, labels=xticks_minor, minor=True, fontsize=7)
        ax.tick_params(axis='x', which='minor', pad=5)
        ax.tick_params(axis='x', which='major', pad=10)

        for i in range(1, int(size_x * 10)):
            ax.plot([i / 10, i / 10], [0, height], linestyle='--', color='black')
        ax.plot(x, z, '+')
        for i, j, text in zip(x, z, [str(i) for i in range(1, count_sensors + 1)]):
            ax.text(i, j - 0.01, text, fontsize=8)
        ax.set_aspect('equal') if height == 0.1 else None
        plt.savefig(f'Отчет {model_name}_{alpha}\\Модель\\Модель {model_name}_{alpha}.png', bbox_inches='tight')

    def generate_envelopes(self, alpha, model_name):
        if model_name[0] == model_name[1]:
            angle_bor = 50
        else:
            angle_bor = 95
        fig, ax = plt.subplots(figsize=(6.49, 5), dpi=200, num=1, clear=True)
        step_x = 20
        step_x_minor = 5
        step_y = 0.4

        step_sens = 100
        count_sensors_plot = len(self.get_face_number(alpha, model_name))

        ax.grid(visible=True, which='minor', color='black', linestyle='--')
        ax.grid(visible=True, which='major', color='black', linewidth=1.5)
        for angle in range(0, angle_bor, 5):
            if not os.path.isdir(
                    f'{os.getcwd()}\\Отчет {model_name}_{alpha}\\Огибающие\\Огибающие {model_name}_{alpha} {angle:02}'):
                os.mkdir(
                    f'{os.getcwd()}\\Отчет {model_name}_{alpha}\\Огибающие\\Огибающие {model_name}_{alpha} {angle:02}')
            pr_coeff_f = np.array(self.get_pressure_coefficients(alpha, model_name, angle)) / 1000
            for q in range(0, count_sensors_plot, 100):
                lines = []
                pr_coeff = pr_coeff_f.T[q:q + step_sens].T

                mean_pr = np.mean(pr_coeff, axis=0).round(4)
                rms_pr = np.array([np.sqrt(i.dot(i) / i.size) for i in pr_coeff.T]).round(4)
                std_pr = np.std(pr_coeff, axis=0).round(4)
                max_pr = np.max(pr_coeff, axis=0).round(4)
                min_pr = np.min(pr_coeff, axis=0).round(4)

                xticks = np.arange(q, q + step_sens + step_x, step_x)
                xticks_minor = np.arange(q, q + step_sens + step_x_minor, step_x_minor)

                ax.set_xticks(xticks, fontsize=10)
                ax.set_xticks(xticks_minor, labels=xticks_minor, minor=True, fontsize=7)

                ax.set_xlim(q, q + step_sens + 1)
                yticks = np.arange(np.min(min_pr) - step_y, np.max(max_pr) + step_y, step_y).round(2)
                ax.set_ylim(np.min(yticks), np.max(yticks))
                ax.set_yticks(yticks)
                ox = [i for i in range(q + 1, q + step_sens + 1)]

                for i, j, c in zip((mean_pr, rms_pr, std_pr, max_pr, min_pr), ('MEAN', 'RMS', 'STD', 'MAX', 'MIN'),
                                   ('b', 'g', 'r', 'c', 'y')):
                    lines.append(ax.plot(ox, i, '-', label=j, linewidth=3, color=c))
                legend = ax.legend(loc='upper right', fontsize=9)
                plt.savefig(
                    f'Отчет {model_name}_{alpha}\\Огибающие\\Огибающие {model_name}_{alpha} {angle:02}\\Огибающие {model_name}_{alpha} {angle:02} {q} - {q + step_sens}.png',
                    bbox_inches='tight')
                for i in lines:
                    i[0].remove()
                legend.remove()

    def get_face_number(self, alpha, model_name):
        if alpha == '4' or alpha == 4:
            self.cursor.execute("""
                        select face_number
                        from experiments_alpha_4
                        where model_name = (%s)
                    """, (model_name,))

        elif alpha == '6' or alpha == 6:
            self.cursor.execute("""
                        select face_number
                        from experiments_alpha_6
                        where model_name = (%s)
                    """, (model_name,))
        self.__connection.commit()
        return self.cursor.fetchall()[0][0]

    def get_uh_average_wind_speed(self, alpha, model_name):
        if alpha == '4' or alpha == 4:
            self.cursor.execute("""
                        select uh_averagewindspeed
                        from experiments_alpha_4
                        where model_name = (%s)
                    """, (model_name,))

        elif alpha == '6' or alpha == 6:
            self.cursor.execute("""
                        select uh_averagewindspeed
                        from experiments_alpha_6
                        where model_name = (%s)
                    """, (model_name,))
        self.__connection.commit()
        return self.cursor.fetchall()[0][0]

    def get_info_sensors(self, alpha, model_name):
        info_sensors = []
        x, z = self.get_coordinates(alpha, model_name)
        for angle in range(0, 50, 5):
            listangle = []
            pr_coeff = np.array(self.get_pressure_coefficients(alpha, model_name, angle)) / 1000
            face_number = self.get_face_number(alpha, model_name)
            breadth, depth, height = int(model_name[0]) / 10, int(model_name[1]) / 10, int(model_name[2]) / 10
            sensors_on_model = len(pr_coeff[0])
            x_new, y_new = self.converter_coordinates(x, depth, breadth, face_number, sensors_on_model)
            mean_pr = np.mean(pr_coeff, axis=0)
            rms_pr = np.array([np.sqrt(i.dot(i) / i.size) for i in pr_coeff.T])
            std_pr = np.std(pr_coeff, axis=0)
            max_pr = np.max(pr_coeff, axis=0)
            min_pr = np.min(pr_coeff, axis=0)
            rach = [Controller.rach(i).round(2) for i in pr_coeff.T]
            _obec_p = [Controller.obes_p(i).round(2) for i in pr_coeff.T]
            _obec_m = [Controller.obes_m(i).round(2) for i in pr_coeff.T]
            for i in range(sensors_on_model):
                row = [i + 1,
                       x_new[i],
                       y_new[i],
                       z[i],
                       mean_pr[i].round(2),
                       rms_pr[i].round(2),
                       std_pr[i].round(2),
                       max_pr[i].round(2),
                       min_pr[i].round(2),
                       rach[i],
                       _obec_p[i],
                       _obec_m[i]
                       ]
                listangle.append(row)
            info_sensors.append(listangle)
        return info_sensors

    def get_sum_coeff(self, alpha, model_name, angle = 0):
        sum_int_x = []
        sum_int_y = []
        pr_coeff = np.array(self.get_pressure_coefficients(alpha, model_name, angle)) / 1000
        count_sensors_on_model = len(pr_coeff[0])
        count_sensors_on_middle = int(model_name[0]) * 5
        count_sensors_on_side = int(model_name[1]) * 5
        count_row = count_sensors_on_model // (2 * (count_sensors_on_middle + count_sensors_on_side))
        for coeff in pr_coeff:

            coeff = np.reshape(coeff, (count_row, -1))
            coeff = np.split(coeff, [count_sensors_on_middle,
                                     count_sensors_on_middle + count_sensors_on_side,
                                     2 * count_sensors_on_middle + count_sensors_on_side,
                                     2 * (count_sensors_on_middle + count_sensors_on_side)
                                     ], axis=1)
            del coeff[4]
            faces_x = []
            faces_y = []
            for face in range(len(coeff)):
                if face in [0, 2]:
                    faces_x.append(np.sum(coeff[face]) / (count_sensors_on_model / 4))
                else:
                    faces_y.append(np.sum(coeff[face]) / (count_sensors_on_model / 4))

            sum_int_x.append((faces_x[0] - faces_x[1]))
            sum_int_y.append((faces_y[0] - faces_y[1]))
        return sum_int_x, sum_int_y

    def get_sum_cmz(self, alpha, model_name, angle = 0):
        sum_cmz = []
        pr_coeff = np.array(self.get_pressure_coefficients(alpha, model_name, angle)) / 1000
        breadth, depth, height = int(model_name[0]) / 10, int(model_name[1]) / 10, int(model_name[2]) / 10
        v2 = breadth
        v3 = breadth + depth
        v4 = 2 * breadth + depth
        mid13_x = breadth / 2
        mid24_x = depth / 2
        count_sensors_on_model = len(pr_coeff[0])
        count_sensors_on_middle = int(model_name[0]) * 5
        count_sensors_on_side = int(model_name[1]) * 5
        count_row = count_sensors_on_model // (2 * (count_sensors_on_middle + count_sensors_on_side))
        x, z = self.get_coordinates(alpha, model_name)

        x = np.reshape(x, (count_row, -1))
        x = np.split(x, [count_sensors_on_middle,
                         count_sensors_on_middle + count_sensors_on_side,
                         2 * count_sensors_on_middle + count_sensors_on_side,
                         2 * (count_sensors_on_middle + count_sensors_on_side)
                         ], axis=1)

        del x[4]
        x[1] -= v2
        x[2] -= v3
        x[3] -= v4
        mx = np.array([
            abs(x[0] - mid13_x),
            abs(x[1] - mid24_x),
            abs(x[2] - mid13_x),
            abs(x[3] - mid24_x),
        ])
        coeffs_norm_13 = [1 if i <= count_sensors_on_middle // 2 else -1 for i in range(count_sensors_on_middle)]
        coeffs_norm_24 = [1 if i <= count_sensors_on_side // 2 else -1 for i in range(count_sensors_on_side)]
        for coeff in pr_coeff:

            coeff = np.reshape(coeff, (count_row, -1))
            coeff = np.split(coeff, [count_sensors_on_middle,
                                     count_sensors_on_middle + count_sensors_on_side,
                                     2 * count_sensors_on_middle + count_sensors_on_side,
                                     2 * (count_sensors_on_middle + count_sensors_on_side)
                                     ], axis=1)
            del coeff[4]
            for i in range(4):
                if i in [0, 2]:
                    coeff[i] *= coeffs_norm_13
                else:
                    coeff[i] *= coeffs_norm_24
            cmz = mx * coeff
            sum_cmz.append(np.sum(cmz))

        return sum_cmz

    def graph_sum_coeff(self, alpha, model_name, sum_cx, sum_cy, sum_cmz):
        info_sum_coeff = []
        fig, ax = plt.subplots(figsize=(6.49, 5), dpi=200, num=1, clear=True)
        ax.grid()
        ax.set_xlim(0, 32.768)
        ax.set_ylabel('Суммарные аэродинамические коэффициенты')
        ax.set_xlabel('Время, с')
        ox = np.linspace(0, 32.768, 32768)
        for i in range(len(sum_cx)):
            plots = []
            anglelist = []
            for data, name in zip((sum_cx[i], sum_cy[i], sum_cmz[i]), ('Cx', 'Cy', 'CMz')):
                plots.append(ax.plot(ox, data, label=name))
                anglelist.append([
                    name,
                    np.mean(data).round(2),
                    Controller.rms(data).round(2),
                    np.std(data).round(2),
                    np.max(data).round(2),
                    np.min(data).round(2),
                    Controller.rach(data).round(2),
                    Controller.obes_p(data).round(2),
                    Controller.obes_m(data).round(2)
                ])
            info_sum_coeff.append(anglelist)
            legend = ax.legend(loc='upper right', fontsize=9)
            plt.savefig(
                f'Отчет {model_name}_{alpha}\\Суммарные аэродинамические коэффициенты в декартовой системе координат\\Суммарные аэродинамические коэффициенты {model_name}_{alpha} {i * 5:02}',
                bbox_inches='tight')
            for i in plots:
                i[0].remove()
            legend.remove()

        return info_sum_coeff

    def spectr_sum_report(self, alpha, model_name, sum_cx, sum_cy, sum_cmz):
        N = 32768
        FD = 1000
        xf = rfftfreq(N, 1 / FD)[1:N // 2]
        count_peaks = 1
        fig, ax = plt.subplots(figsize=(9, 5), dpi=200, num=1, clear=True)
        ax.set_xlabel('Частота f, Гц')
        ax.set_ylabel('S(C)')

        ax.grid()
        for i in range(len(sum_cx)):
            ax.set_xlim([0, 15])
            ax.set_xscale('linear')
            ax.set_yscale('linear')
            plots = []
            annotates = []
            for data, name in zip((sum_cx[i], sum_cy[i], sum_cmz[i]), ('Cx', 'Cy', 'CMz')):
                spectr = (1 / N) * (np.abs(fft(data)))[1:N // 2]
                plots.append(ax.plot(xf, spectr, antialiased=True, label=name))
                peaks = np.sort(spectr)[::-1][:count_peaks]
                for peak in range(count_peaks):
                    x = xf[np.where(spectr == peaks[peak])].round(3)[0]
                    y = peaks[peak]
                    plots.append(
                    ax.plot([], [], label=f'{x}', color=plots[-1][0]._color))

                    #annotates.append(ax.annotate(x, xy=(x, y + y * 0.03), fontsize=14))

            legend = ax.legend(loc='upper right', fontsize=9)

            plt.savefig(
                f'Отчет {model_name}_{alpha}\\Спектры cуммарных значений аэродинамических коэффициентов\\Линейная шкала\\Спектр cуммарных значений аэродинамических коэффициентов {model_name}_{alpha} {i * 5:02}.png',
                bbox_inches='tight')

            ax.set_xlim([10 ** (-2), 10 ** (3)])
            ax.set_xscale('log')
            ax.set_yscale('log')

            plt.savefig(
                f'Отчет {model_name}_{alpha}\\Спектры cуммарных значений аэродинамических коэффициентов\\Логарифмическая шкала\\Спектр cуммарных значений аэродинамических коэффициентов {model_name}_{alpha} {i * 5:02}.png',
                bbox_inches='tight')
            for i in plots:
                i[0].remove()
            for i in annotates:
                i.remove()
            legend.remove()

    def spect_sum(self, alpha, model_name, angle, parameters):
        labels = parameters['labels']
        data_sum = parameters['data']
        data_spectr = []
        N = len(data_sum[0])
        for i in data_sum:
            data_spectr.append((1 / N) * (np.abs(fft(i)))[1:N // 2])

        FD = 1000
        xf = rfftfreq(N, 1 / FD)[1:N // 2]

        border = 500

        fig, ax = plt.subplots(figsize=(6.49, 5), dpi=200, num=1, clear=True)
        ax.set_xlabel('Частота f, Гц')
        ax.set_ylabel('S(C)')
        for data, name in zip(data_spectr, labels):
            ax.plot(xf[:border], data[:border], antialiased=True, label=name)
            peaks = np.sort(data[:border])[::-1]
            x = xf[:border][np.where(data[:border] == peaks[0])].round(3)[0]
            y = peaks[0]
            ax.annotate(x, xy=(x, y))
        ax.legend()
        ax.grid()
        plt.savefig(
            f'Отчет {model_name}_{alpha}\\Спектры суммарных сил\\Спектр суммарных сил {model_name}_{alpha}_{angle}.png',
            bbox_inches='tight')

    def generate_spectr_sum_pres(self, alpha, model_name, sum_cx, sum_cy, sum_cmz):
        for i in range(len(sum_cx)):
            self.spect_sum(alpha, model_name, i * 5, {'labels': ('Cx', 'CY', 'CMz'),
                                                      'data': (sum_cx[i], sum_cy[i], sum_cmz[i])})

    def get_st(self, alpha, model_name, count_sensors):
        fi = []
        speed = self.get_uh_average_wind_speed(alpha, model_name)
        size = float(model_name[0]) / 10
        N = 32768
        FD = 1000
        border = 500
        count_peaks = 1
        for angle in range(0, 50, 5):
            filist = []

            pr_coeff = np.array(self.get_pressure_coefficients(alpha, model_name, angle)).T / 1000

            fig, ax = plt.subplots(figsize=(6.49, 5), dpi=200, num=1, clear=True)
            for i in range(count_sensors):

                yf = (1 / N) * (np.abs(fft(pr_coeff[i])))[1:N // 2]
                xf = rfftfreq(N, 1 / FD)[1:N // 2]
                ax.plot(xf[:border], yf[:border], antialiased=True, label=i + 1)

                peaks = np.sort(yf[:border])[::-1][:count_peaks]

                for peak in range(count_peaks):
                    x = xf[:border][np.where(yf[:border] == peaks[peak])].round(3)[0]
                    filist.append(x)
            fi.append((np.array(filist) * size / speed).round(4))
            # ax.legend()
            ax.grid()
            plt.savefig(f'Отчет {model_name}_{alpha}\\числа Струхаля для датчиков {model_name}_{alpha}_{angle}.png')
        return np.array(fi).T

    def get_info_sens_st(self, alpha, model_name):
        breadth, depth = int(model_name[0]) / 10, int(model_name[1]) / 10
        info_sens_st = []
        x, z = self.get_coordinates(alpha, model_name)
        face_number = self.get_face_number(alpha, model_name)
        x, y = self.converter_coordinates(x, depth, breadth, face_number, len(face_number))
        st = self.get_st(alpha, model_name, len(face_number))
        for i in range(len(face_number)):
            info_sens_st.append([i + 1, x[i], y[i], z[i], *st[i]])
        return info_sens_st

    def generate_isofields(self, alpha, model_name):
        coords = self.get_coordinates(alpha, model_name)
        for angle in range(0, 50, 5):
            coefs = np.array(self.get_pressure_coefficients(alpha, model_name, angle)) / 1000
            for mode in ('max', 'mean', 'min', 'std'):
                Artist.isofield(mode, coefs, coords, alpha, model_name, angle)
                Artist.disc_isofield(mode, coefs, coords, alpha, model_name, angle)

    @staticmethod
    def rms(data):
        return np.sqrt(np.array(data).dot(np.array(data)) / np.array(data).size).round(2)

    @staticmethod
    def rach(data):
        return np.max([np.abs(np.min(data)), np.abs(np.max(data))]).round(2)

    @staticmethod
    def obes_p(data):
        return (np.abs(np.max(data) - np.mean(data)) / np.std(data)).round(2)

    @staticmethod
    def obes_m(data):
        return (np.abs(np.min(data) - np.mean(data)) / np.std(data)).round(2)

    @staticmethod
    def polar_plot(alpha, model_name, title, cx, cy, cmz):
        angles = np.array([angle for angle in range(0, 365, 5)]) * np.pi / 180.0
        fig, ax = plt.subplots(figsize=(6.49, 5), dpi=200, num=1, clear=True, subplot_kw={'projection': 'polar'})

        for i, j, label in zip((cy, cx, cmz), (cx, cy, cmz), ('CX', 'CY', 'CMz')):
            a = np.array(i)
            b = np.append(a, np.flip(j)[1:])
            c = np.append(b, np.flip(b)[1:])
            d = np.append(c, np.flip(c)[1:])
            ax.plot(angles, d, label=label)

        ax.set_theta_direction(-1)
        ax.set_theta_zero_location('N')
        ax.set_thetagrids([i for i in range(0, 360, 15)])
        ax.legend(loc='upper right', fontsize=9)
        ax.set_title(title)
        plt.savefig(
            f'Отчет {model_name}_{alpha}\\Суммарные аэродинамические коэффициенты в полярной системе координат\\Суммарные аэродинамические коэффициенты в полярной системе координат {model_name}_{alpha} {title}.png',
            bbox_inches='tight')

    @staticmethod
    def sum_polar_plot(alpha, model_name, sum_cx, sum_cy, sum_cmz):
        mods = {
            'MEAN': np.mean,
            'RMS': Controller.rms,
            'STD': np.std,
            'MAX': np.max,
            'MIN': np.min,
            'РАСЧЕТНОЕ': Controller.rach,
            'ОБЕСП+': Controller.obes_m,
            'ОБЕСП-': Controller.obes_p
        }
        for mode in mods.keys():
            cx = np.array([mods[mode](i) for i in sum_cx])
            cy = np.array([mods[mode](i) for i in sum_cy])
            cmz = np.array([mods[mode](i) for i in sum_cmz])
            Controller.polar_plot(alpha, model_name, mode, cx, cy, cmz)

    def generate_welch_graphs(self, alpha, model_name, sum_cx, sum_cy, sum_cmz):
        numSt = []
        speed = self.get_uh_average_wind_speed(alpha, model_name)
        size = float(model_name[0]) / 10
        fig, ax = plt.subplots(figsize=(6.49, 5), dpi=200, num=1, clear=True)
        ax.set_xlim([0, 15])
        ax.grid()
        ax.set_xlabel('Sh')
        ax.set_ylabel('PSD, V**2/Hz')
        frequency = [i for i in range(1, 16)]
        for i in range(len(sum_cx)):
            ax.set_xlim([0, 15])
            ax.set_xscale('linear')
            ax.set_yscale('linear')
            iterSt = [i * 5]
            plots = []
            annotates = []
            for data, name in zip((sum_cx[i], sum_cy[i], sum_cmz[i]), ('Cx', 'Cy', 'CMz')):
                temp, psd = welch(data, fs=1000, nperseg=int(32768 / 5))
                plots.append(ax.plot(temp, psd, label=name))
                peak = np.max(psd)
                x = temp[np.where(psd == peak)]
                y = peak
                plots.append(ax.plot([], [], label=f'Sh {np.array(x * size / speed).round(2)[0]}', color=plots[-1][0]._color))
                plots.append(ax.plot([], [], label=f'Частота {np.array(x).round(2)[0]}', color=plots[-1][0]._color))
                #annotates.append(ax.annotate(np.array(x * size / speed).round(4)[0], xy=(x, y)))
                iterSt.append(np.array(x * size / speed).round(4)[0])
            numSt.append(iterSt)
            ax.set_xticks(frequency, labels=[np.array(i * size / speed).round(3) for i in frequency], fontsize=7)



            legend = ax.legend(loc='upper right', fontsize=9)

            plt.savefig(
                f'Отчет {model_name}_{alpha}\\Спектральная плотность мощности\\Линейная шкала\\Спектральная плотность мощности {model_name}_{alpha} {i * 5:02}.png',
                bbox_inches='tight')

            ax.set_xlim([10 ** (-2), 10 ** (3)])
            ax.set_xscale('log')
            ax.set_yscale('log')

            plt.savefig(
                f'Отчет {model_name}_{alpha}\\Спектральная плотность мощности\\Логарифмическая шкала\\Спектральная плотность мощности {model_name}_{alpha} {i * 5:02}.png',
                bbox_inches='tight')

            for i in plots:
                i[0].remove()
            for i in annotates:
                i.remove()
            legend.remove()
        return numSt

    def generate_model_polar(self, alpha, model_name):
        b_scale, d_scale, h_scale = int(model_name[0]), int(model_name[1]), int(model_name[2])
        rect = 1, 1, 1, 1
        fig = plt.figure(figsize=(6.49 / 2, 5), dpi=200, num=1, clear=True)
        ax = fig.add_axes(rect)
        polar = fig.add_axes(rect, polar=True)
        polar.patch.set_alpha(0)
        polar.set_theta_zero_location('N')
        polar.set_theta_direction(-1)
        polar.set_thetagrids([i for i in range(0, 360, 15)], fontsize=7)
        polar.set_yticks([0, 1, 2, 3])
        polar.set_yticklabels([0, 1, 2, 3], visible=False)
        polar.set_autoscale_on(False)
        angles = np.array([angle for angle in range(0, 365, 5)]) * np.pi / 180.0
        polar.annotate("", xy=(angles[0], 2), xytext=(0, 0),
                       arrowprops=dict(arrowstyle="->",
                                       linewidth=2.5))
        polar.annotate("", xy=(angles[18], 2), xytext=(0, 0),
                       arrowprops=dict(arrowstyle="->",
                                       linewidth=2.5))
        polar.annotate("y", xy=(angles[0], 2))
        polar.annotate("x", xy=(angles[18], 2))
        ax.set_visible(True)
        ax.set_autoscale_on(False)
        ax.set_aspect(1)
        ax.axis('off')
        dx = dy = (0.7 - 0.3) / 3
        b = dx * b_scale
        d = dy * d_scale
        mid = 0.5
        x0 = mid - b / 2
        x1 = mid + b / 2
        y0 = mid - d / 2
        y1 = mid + d / 2
        ax.fill_between([x0, x1], [y0, y0], [y1, y1], color='grey', alpha=0.5)
        plt.savefig(
            f'Отчет {model_name}_{alpha}\\Модель\\Система координат и расчетные направления ветрового потока {model_name}_{alpha}.png',
            bbox_inches='tight')

    def generate_model_cube(self, alpha, model_name):
        b_scale, d_scale, h_scale = int(model_name[0]), int(model_name[1]), int(model_name[2])
        fig = plt.figure(figsize=(6.4/2, 5), dpi=200, num=1, clear=True)
        ax = fig.add_subplot(111, projection='3d', box_aspect=(b_scale + 1, d_scale + 1, h_scale + 1))
        count_nodes = 2
        x = np.linspace(1, b_scale + 1, count_nodes)
        y = np.linspace(1, d_scale + 1, count_nodes)
        z = np.linspace(0, h_scale, count_nodes)

        ones2d = np.ones((count_nodes, count_nodes))
        ones1d = np.ones(count_nodes)
        zeros2d = np.zeros((count_nodes, count_nodes))

        ax.set_xlim([1, b_scale + 1])
        ax.set_ylim([1, d_scale + 1])
        ax.set_zlim([0, h_scale])

        X_top, Y_top = np.meshgrid(x, y)
        Z_top = zeros2d + h_scale
        ax.plot_surface(X_top, Y_top, Z_top, color='grey')

        X_face, Y_face = np.meshgrid(x, ones1d)
        Z_face = (zeros2d + z).T
        ax.plot_surface(X_face, Y_face, Z_face, color='grey')

        Y_side, X_side = np.meshgrid(y, ones1d + b_scale)
        ax.plot_surface(X_side, Y_side, Z_face, color='grey')

        ax.axis('off')
        # ax.set_xlabel('Depth')
        # ax.set_ylabel('Breadth')
        # ax.set_zlabel('Height')

        plt.savefig(f'Отчет {model_name}_{alpha}\\Модель\\3D Модель {model_name}_{alpha}.png', bbox_inches='tight')


    def generate_folder_report(self, alpha, model_name):
        if not os.path.isdir(f'{os.getcwd()}\\Отчет {model_name}_{alpha}'):
            os.mkdir(f'{os.getcwd()}\\Отчет {model_name}_{alpha}')
        folders = (
            'Модель',
            'Огибающие',
            'Изополя ветровых нагрузок и воздействий',
            'Изополя ветровых нагрузок и воздействий\\Непрерывные',
            'Изополя ветровых нагрузок и воздействий\\Дискретные',
            'Суммарные аэродинамические коэффициенты в полярной системе координат',
            'Суммарные аэродинамические коэффициенты в декартовой системе координат',
            'Спектры cуммарных значений аэродинамических коэффициентов',
            'Спектры cуммарных значений аэродинамических коэффициентов\\Линейная шкала',
            'Спектры cуммарных значений аэродинамических коэффициентов\\Логарифмическая шкала',
            'Спектральная плотность мощности',
            'Спектральная плотность мощности\\Линейная шкала',
            'Спектральная плотность мощности\\Логарифмическая шкала',
        )
        for folder in folders:
            if not os.path.isdir(f'{os.getcwd()}\\Отчет {model_name}_{alpha}\\{folder}'):
                os.mkdir(f'{os.getcwd()}\\Отчет {model_name}_{alpha}\\{folder}')

    def generate_report(self, alpha, model_name):
        breadth, depth, height = int(model_name[0]) / 10, int(model_name[1]) / 10, int(model_name[2]) / 10
        self.generate_folder_report(alpha, model_name)
        sum_cmz_list = []
        sum_x_list = []
        sum_y_list = []

        for angle in range(0, 50, 5):
            print(angle)
            sum_cmz_list.append(self.get_sum_cmz(alpha, model_name, angle))
            sum_int_x, sum_int_y = self.get_sum_coeff(alpha, model_name, angle)
            sum_x_list.append(sum_int_x)
            sum_y_list.append(sum_int_y)

        print("Работа с данными")
        Controller.sum_polar_plot(alpha, model_name, sum_x_list, sum_y_list, sum_cmz_list)
        print('1\\7')
        self.generate_model_pic(alpha, model_name)
        self.generate_model_polar(alpha, model_name)
        self.generate_model_cube(alpha, model_name)
        print('2\\7')
        self.generate_envelopes(alpha, model_name)
        print('3\\7')
        self.generate_isofields(alpha, model_name)
        print('4\\7')
        numSt = self.generate_welch_graphs(alpha, model_name, sum_x_list, sum_y_list, sum_cmz_list)
        print('5\\7')
        self.spectr_sum_report(alpha, model_name, sum_x_list, sum_y_list, sum_cmz_list)
        print('6\\7')
        info_sensors = self.get_info_sensors(alpha, model_name)
        info_sum_coeff = self.graph_sum_coeff(alpha, model_name, sum_x_list, sum_y_list, sum_cmz_list)
        ####self.generate_spectr_sum_pres(alpha, model_name, sum_x_list, sum_y_list, sum_cmz_list)
        print('7\\7')
        print("Формирование отчета")

        counter_plots = 1
        counter_tables = 1
        doc = Document()
        style = doc.styles['Normal']
        style.font.size = Pt(14)
        style.font.name = 'Times New Roman'
        section = doc.sections[0]
        section.left_margin = Mm(30)
        section.right_margin = Mm(15)
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        ####################### Шапка #######################
        title = doc.add_heading().add_run(f'Отчет по зданию {breadth}x{depth}x{height}')
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.font.size = Pt(24)
        title.bold = True
        for i in ('Параметры ветрового воздействия:',
                  'Ветровой район: None',
                  'Тип местности: None'
                  ):
            doc.add_paragraph().add_run(i)
        p = doc.add_paragraph()

        run = p.add_run()
        run.add_picture(f'Отчет {model_name}_{alpha}\\Модель\\3D Модель {model_name}_{alpha}.png', width=Mm(82.5))
        run.add_picture(
            f'Отчет {model_name}_{alpha}\\Модель\\Система координат и расчетные направления ветрового потока {model_name}_{alpha}.png',
            width=Mm(82.5))
        doc.add_paragraph().add_run(
            f'Рисунок {counter_plots}. Геометрические размеры и система координат направления ветровых потоков')
        counter_plots += 1
        ####################### 1.Геометрические размеры здания #######################
        doc.add_heading().add_run('1. Геометрические размеры здания').font.size = Pt(20)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        header_model = (
            'Геометрический размер',
            'Значение, м'
        )
        table_model = doc.add_table(rows=1, cols=len(header_model))
        table_model.style = 'Table Grid'
        hdr_cells = table_model.rows[0].cells
        for i in range(len(header_model)):
            hdr_cells[i].add_paragraph().add_run(header_model[i])
        for i, j in zip((breadth, depth, height), ('Ширина:', 'Глубина:', 'Высота:')):
            row_cells = table_model.add_row().cells
            row_cells[0].add_paragraph().add_run(j)
            row_cells[1].add_paragraph().add_run(str(i))
        doc.add_picture(f'Отчет {model_name}_{alpha}\\Модель\\Модель {model_name}_{alpha}.png')
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph().add_run(f'Рисунок {counter_plots}. Система датчиков мониторинга')
        counter_plots += 1
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        print('1\\7')
        doc.add_page_break()
        ####################### 2. Статистика по датчиках. Максимумы и огибающие #######################
        doc.add_heading().add_run('2. Статистика по датчиках. Максимумы и огибающие').font.size = Pt(20)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for angle in range(0, 50, 5):
            envelopes = glob.glob(
                f'Отчет {model_name}_{alpha}\\Огибающие\\Огибающие {model_name}_{alpha} {angle:02}\\Огибающие *.png')
            for i in envelopes:
                doc.add_picture(i, height=Mm(80))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph().add_run(
                    f'Рисунок {counter_plots}. Огибающая ветрового давления для здания {breadth}x{depth}x{height} при угле {angle:02}º датчики {i[i[:i.rfind("-")-1].rfind(" ")+1:i.rfind(".")]}')
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                counter_plots += 1
        print('2\\7')
        doc.add_page_break()
        ####################### 3. Изополя ветровых нагрузок и воздействий #######################
        doc.add_heading().add_run('3. Изополя ветровых нагрузок и воздействий').font.size = Pt(20)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_heading(level=2).add_run('3.1 Непрерывные изополя').font.size = Pt(16)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        isofields = glob.glob(f'Отчет {model_name}_{alpha}\\Изополя ветровых нагрузок и воздействий\\Непрерывные\\Изополя *.png')
        for i in isofields:
            doc.add_picture(i)
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            angle = i[i.rfind('_') + 1:i.rfind(' ')]
            mode = i[i.rfind(" ") + 1:i.rfind(".")]
            doc.add_paragraph().add_run(
                f'Рисунок {counter_plots}. Непрерывные изополя {mode} для здания {breadth}x{depth}x{height} при угле {angle}º')
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            counter_plots += 1
        doc.add_heading(level=2).add_run('3.2 Дискретные изополя').font.size = Pt(16)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        isofields_disc = glob.glob(f'Отчет {model_name}_{alpha}\\Изополя ветровых нагрузок и воздействий\\Дискретные\\Изополя *.png')
        for i in isofields_disc:
            doc.add_picture(i)
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            angle = i[i.rfind('_') + 1:i.rfind(' ')]
            mode = i[i.rfind(" ") + 1:i.rfind(".")]
            doc.add_paragraph().add_run(
                f'Рисунок {counter_plots}. Дискретные изополя {mode} для здания {breadth}x{depth}x{height} при угле {angle}º')
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            counter_plots += 1
        print('3\\7')

        doc.add_page_break()
        ####################### 4. Статистика по датчикам в табличном виде #######################
        doc.add_heading().add_run('4. Статистика по датчикам в табличном виде ').font.size = Pt(20)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_sensors = (
            'ДАТЧИК',
            'X(мм)',
            'Y(мм)',
            'Z(мм)',
            'MEAN',
            'RMS',
            'STD',
            'MAX',
            'MIN',
            'РАСЧЕТНОЕ',
            'ОБЕСП+',
            'ОБЕСП-'
        )
        for angle in range(0, 50, 5):
            doc.add_paragraph().add_run(
                f'\nТаблица {counter_tables}. Аэродинамический коэффициент в датчиках для здания {breadth}x{depth}x{height} при угле {angle}º')
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            counter_tables += 1
            table_sensors = doc.add_table(rows=1, cols=len(header_sensors))
            table_sensors.style = 'Table Grid'
            hdr_cells = table_sensors.rows[0].cells
            for i in range(len(header_sensors)):
                hdr_cells[i].add_paragraph().add_run(header_sensors[i]).font.size = Pt(8)

            for rec in info_sensors[angle % 5]:
                row_cells = table_sensors.add_row().cells
                for i in range(len(rec)):
                    row_cells[i].add_paragraph().add_run(str(rec[i])).font.size = Pt(12)

        print('4\\7')
        doc.add_page_break()
        ####################### 5. Суммарные значения аэродинамических коэффициентов #######################
        doc.add_page_break()
        doc.add_heading().add_run('5. Суммарные значения аэродинамических коэффициентов').font.size = Pt(20)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_sum = (
            'СИЛА',
            'MEAN',
            'RMS',
            'STD',
            'MAX',
            'MIN',
            'РАСЧЕТНОЕ',
            'ОБЕСП+',
            'ОБЕСП-'
        )
        for angle in range(0, 50, 5):
            doc.add_picture(
                f'Отчет {model_name}_{alpha}\\Суммарные аэродинамические коэффициенты в декартовой системе координат\\Суммарные аэродинамические коэффициенты {model_name}_{alpha} {angle:02}.png')
            doc.add_paragraph().add_run(
                f'Рисунок {counter_plots}. Суммарные аэродинамические коэффициенты для здания {breadth}x{depth}x{height} при угле {angle}º')
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            counter_plots += 1
            doc.add_paragraph().add_run(
                f'Таблица {counter_tables}. Суммарные аэродинамические коэффициенты для здания {breadth}x{depth}x{height} при угле {angle}º')
            counter_tables += 1
            table_sum = doc.add_table(rows=1, cols=len(header_sum))
            table_sum.style = 'Table Grid'
            hdr_cells = table_sum.rows[0].cells
            for i in range(len(header_sum)):
                hdr_cells[i].add_paragraph().add_run(header_sum[i]).font.size = Pt(8)

            for rec in info_sum_coeff[angle % 5]:
                row_cells = table_sum.add_row().cells
                for i in range(len(rec)):
                    row_cells[i].add_paragraph().add_run(str(rec[i])).font.size = Pt(12)

            doc.add_page_break()

        polar_graphs = glob.glob(
            f'Отчет {model_name}_{alpha}\\Суммарные аэродинамические коэффициенты в полярной системе координат\\Суммарные аэродинамические коэффициенты в полярной системе координат *.png')
        for i in polar_graphs:
            doc.add_picture(i)
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            mode = i[i.rfind(" ") + 1:i.rfind(".")]
            doc.add_paragraph().add_run(
                f'Рисунок {counter_plots}. Суммарные аэродинамические коэффициенты в полярной системе координат {mode} для здания {breadth}x{depth}x{height}')
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            counter_plots += 1
        print('5\\7')
        doc.add_page_break()
        ####################### 6. Спектры cуммарных значений аэродинамических коэффициентов #######################
        doc.add_heading().add_run('6. Спектры cуммарных значений аэродинамических коэффициентов').font.size = Pt(20)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        spectrs = glob.glob(
            f'Отчет {model_name}_{alpha}\\Спектры cуммарных значений аэродинамических коэффициентов\\Логарифмическая шкала\\Спектр cуммарных значений аэродинамических коэффициентов *.png')
        for i in spectrs:
            doc.add_picture(i)
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            angle = i[i.rfind(" ") + 1:i.rfind(".")]
            doc.add_paragraph().add_run(
                f'Рисунок {counter_plots}. Спектр cуммарных значений аэродинамических коэффициентов для здания {breadth}x{depth}x{height} при угле {angle}º')
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            counter_plots += 1
        print('6\\7')
        doc.add_page_break()
        ####################### 7. Числа Струхаля #######################
        doc.add_paragraph().add_run('7. Числа Струхаля').font.size = Pt(20)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_st = (
            'fi',
            'St(CX)',
            'St(CY)',
            'St(CMz)'
        )

        spectrs_st = glob.glob(
            f'Отчет {model_name}_{alpha}\\Спектральная плотность мощности\\Логарифмическая шкала\\Спектральная плотность мощности *.png')
        for i in spectrs_st:
            doc.add_picture(i)
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            angle = i[i.rfind(" ") + 1:i.rfind(".")]
            doc.add_paragraph().add_run(
                f'Рисунок {counter_plots}. Спектральная плотность мощности для здания {breadth}x{depth}x{height} при угле {angle:02}º')
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            counter_plots += 1
        table_st = doc.add_table(rows=1, cols=len(header_st))
        table_st.style = 'Table Grid'
        hdr_cells = table_st.rows[0].cells
        for i in range(len(header_st)):
            hdr_cells[i].add_paragraph().add_run(header_st[i]).font.size = Pt(8)

        for rec in numSt:
            row_cells = table_st.add_row().cells
            for i in range(len(rec)):
                row_cells[i].add_paragraph().add_run(str(rec[i])).font.size = Pt(12)

        doc.add_paragraph().add_run(
            f'Таблица {counter_tables}. Числа Струхаля для здания {breadth}x{depth}x{height}')
        counter_tables += 1
        print('7\\7')
        doc.save(f'Отчет {model_name}_{alpha}\\Отчет по зданию {breadth}x{depth}x{height}.docx')


if __name__ == '__main__':
    control = Controller()
    # пароль 08101430
    control.connect(database='tpu', password='2325070307')

    control.generate_report('6', '111')
    control.disconnect()
